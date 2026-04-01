"""
generate_sample_data.py - Create synthetic university documents for testing.

Generates:
  data/handbook_2024.pdf
  data/circular_exam_rules_2023.pdf
  data/circular_exam_rules_2022.pdf   ← outdated duplicate
  data/cse_department_policy.docx
  data/internal_staff_policy.html
"""

from __future__ import annotations
import os
from pathlib import Path

DATA_DIR = Path(__file__).parent / "data"
DATA_DIR.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# PDF generation using reportlab (or plain text fallback)
# ---------------------------------------------------------------------------

def _write_pdf(path: Path, title: str, sections: list[tuple[str, str]]) -> None:
    """Write a multi-section PDF. Falls back to .txt if reportlab not installed."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm

        doc = SimpleDocTemplate(str(path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = [Paragraph(title, styles["Title"]), Spacer(1, 0.5 * cm)]

        for heading, body in sections:
            story.append(Paragraph(heading, styles["Heading1"]))
            story.append(Spacer(1, 0.2 * cm))
            for para in body.strip().split("\n\n"):
                story.append(Paragraph(para.replace("\n", " "), styles["BodyText"]))
                story.append(Spacer(1, 0.15 * cm))

        doc.build(story)
        print(f"  ✓ PDF: {path.name}")

    except ImportError:
        # Fallback: write a plain text file that PyMuPDF will still load
        txt_path = path.with_suffix(".txt")
        with open(txt_path, "w") as f:
            f.write(f"{title}\n{'=' * len(title)}\n\n")
            for heading, body in sections:
                f.write(f"\n{heading}\n{'-' * len(heading)}\n{body}\n")
        # Rename to .pdf (PyMuPDF can't open it but our fallback loader handles .txt)
        # Instead, write it as a PDF using a minimal PDF structure:
        _write_minimal_pdf(path, title, sections)
        print(f"  ✓ PDF (minimal): {path.name}")


def _write_minimal_pdf(path: Path, title: str, sections: list[tuple[str, str]]) -> None:
    """Write a minimal valid PDF with text content."""
    lines = [title, "=" * 60, ""]
    for heading, body in sections:
        lines.append(heading)
        lines.append("-" * 40)
        lines.append(body)
        lines.append("")

    full_text = "\n".join(lines)
    # Build a minimal PDF manually
    pdf_bytes = _build_simple_pdf(full_text)
    path.write_bytes(pdf_bytes)


def _build_simple_pdf(text: str) -> bytes:
    """Build a minimal valid 1-page PDF containing the given text."""
    # Escape special PDF chars
    safe = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    # Split into lines of ~80 chars each for PDF stream
    raw_lines = safe.split("\n")
    pdf_lines = []
    for line in raw_lines:
        # Wrap at 80 chars
        while len(line) > 80:
            pdf_lines.append(line[:80])
            line = line[80:]
        pdf_lines.append(line)

    # Build BT/ET text block (12pt Helvetica, line spacing 14)
    text_ops = []
    y = 780
    for ln in pdf_lines[:55]:  # Max ~55 lines on A4
        text_ops.append(f"BT /F1 10 Tf {50} {y} Td ({ln}) Tj ET")
        y -= 14
        if y < 50:
            break

    stream = "\n".join(text_ops)
    stream_bytes = stream.encode("latin-1", errors="replace")

    objects = []

    def obj(n: int, content: str) -> None:
        objects.append((n, content))

    obj(1, "<< /Type /Catalog /Pages 2 0 R >>")
    obj(2, "<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    obj(
        3,
        (
            "<< /Type /Page /Parent 2 0 R "
            "/MediaBox [0 0 595 842] "
            "/Contents 4 0 R "
            "/Resources << /Font << /F1 5 0 R >> >> >>"
        ),
    )
    obj(4, f"<< /Length {len(stream_bytes)} >>\nstream\n{stream}\nendstream")
    obj(
        5,
        (
            "<< /Type /Font /Subtype /Type1 "
            "/BaseFont /Helvetica "
            "/Encoding /WinAnsiEncoding >>"
        ),
    )

    # Assemble
    out = [b"%PDF-1.4\n"]
    offsets = {}
    for n, content in objects:
        offsets[n] = sum(len(x) for x in out)
        out.append(f"{n} 0 obj\n{content}\nendobj\n".encode("latin-1", errors="replace"))

    xref_offset = sum(len(x) for x in out)
    xref = f"xref\n0 {len(objects)+1}\n0000000000 65535 f \n"
    for n, _ in objects:
        xref += f"{offsets[n]:010d} 00000 n \n"
    out.append(xref.encode())
    trailer = (
        f"trailer\n<< /Size {len(objects)+1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    )
    out.append(trailer.encode())
    return b"".join(out)


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def _write_docx(path: Path, title: str, sections: list[tuple[str, str]]) -> None:
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(title, 0)

        for heading, body in sections:
            doc.add_heading(heading, level=1)
            for para in body.strip().split("\n\n"):
                doc.add_paragraph(para.replace("\n", " "))

        doc.save(str(path))
        print(f"  ✓ DOCX: {path.name}")
    except ImportError:
        # Plain text fallback
        with open(path.with_suffix(".txt"), "w") as f:
            f.write(f"{title}\n\n")
            for h, b in sections:
                f.write(f"{h}\n{b}\n\n")
        print(f"  ✓ TXT fallback for DOCX: {path.name}")


# ---------------------------------------------------------------------------
# HTML generation
# ---------------------------------------------------------------------------

def _write_html(path: Path, title: str, sections: list[tuple[str, str]]) -> None:
    html = f"<!DOCTYPE html><html><head><title>{title}</title></head><body>\n"
    html += f"<h1>{title}</h1>\n"
    for heading, body in sections:
        html += f"<h2>{heading}</h2>\n"
        for para in body.strip().split("\n\n"):
            html += f"<p>{para.replace(chr(10), ' ')}</p>\n"
    html += "</body></html>"
    path.write_text(html, encoding="utf-8")
    print(f"  ✓ HTML: {path.name}")


# ---------------------------------------------------------------------------
# Document content definitions
# ---------------------------------------------------------------------------

HANDBOOK_2024 = (
    "University Student Handbook 2024",
    [
        (
            "1. Introduction",
            "Welcome to the University. This handbook outlines the rules, regulations, and "
            "resources available to all enrolled students for the academic year 2024.\n\n"
            "The University is committed to providing a safe and inclusive learning environment. "
            "All students are expected to uphold the values of academic integrity, respect, and "
            "responsibility.",
        ),
        (
            "2. Academic Policies",
            "Students must maintain a minimum GPA of 2.0 to remain in good academic standing. "
            "Any student falling below this threshold will be placed on academic probation for "
            "one semester.\n\n"
            "Attendance policy: Students must attend at least 75% of scheduled classes. "
            "Exceeding the absence limit may result in grade penalties or course failure.",
        ),
        (
            "3. Examination Rules",
            "All examinations are conducted under strict invigilation. Students must carry "
            "their university ID cards to examination halls.\n\n"
            "Mobile phones and electronic devices are strictly prohibited during exams. "
            "Any violation will result in immediate disqualification from the examination.",
        ),
        (
            "4. Library and Resources",
            "The university library is open Monday to Saturday, 8:00 AM to 10:00 PM. "
            "Students may borrow up to 5 books at a time for a period of 14 days.\n\n"
            "Digital library access is available 24/7 through the student portal. "
            "Remote access requires valid student credentials.",
        ),
        (
            "5. Grievance Redressal",
            "Students with grievances may approach the Student Affairs Office located in "
            "Block A, Room 105. Grievances must be submitted in writing within 30 days of "
            "the incident.\n\n"
            "An independent Grievance Committee reviews all submissions within 15 working days. "
            "The committee's decision is binding unless escalated to the Academic Council.",
        ),
        (
            "6. Fee Structure and Scholarships",
            "Tuition fees for the academic year 2024 are as follows: B.Tech Rs 85,000/year, "
            "M.Tech Rs 65,000/year, MBA Rs 1,10,000/year.\n\n"
            "Merit scholarships are awarded to students ranking in the top 10% of their batch. "
            "Need-based financial aid is available through the Welfare Fund.",
        ),
    ],
)

CIRCULAR_EXAM_2023 = (
    "Circular: Examination Rules and Regulations – 2023",
    [
        (
            "Subject: End Semester Examination Guidelines",
            "This circular is issued to inform all students and faculty regarding the "
            "updated examination procedures effective from November 2023 examination cycle.\n\n"
            "All departments are advised to communicate these guidelines to students at least "
            "two weeks before the commencement of examinations.",
        ),
        (
            "1. Hall Ticket",
            "Hall tickets will be issued digitally through the student portal from "
            "15 October 2023. Students must download and print the hall ticket.\n\n"
            "Physical hall tickets will not be issued at the examination centre. "
            "Students without a valid hall ticket will not be permitted to appear.",
        ),
        (
            "2. Seating Arrangement",
            "Seating arrangements will be displayed on the university noticeboard and "
            "the official website 48 hours before the commencement of each examination.\n\n"
            "Students must sit in the allotted seat only. Changing seats without the "
            "invigilator's permission is a punishable offence.",
        ),
        (
            "3. Prohibited Items",
            "The following items are strictly prohibited inside the examination hall: "
            "mobile phones, smartwatches, earphones, Bluetooth devices, electronic "
            "calculators (unless permitted), and any written material.\n\n"
            "Students found with prohibited items will face disciplinary action as per "
            "the University Examination Ordinance 2020.",
        ),
        (
            "4. Answer Booklet",
            "Students must write their register number and subject code clearly on the "
            "answer booklet. Use blue or black ballpoint pen only.\n\n"
            "Additional sheets will be provided on request. All answer sheets must be "
            "submitted to the invigilator before leaving the hall.",
        ),
    ],
)

CIRCULAR_EXAM_2022 = (
    "Circular: Examination Rules and Regulations – 2022 (OUTDATED)",
    [
        (
            "Subject: End Semester Examination Guidelines",
            "This circular is issued to inform all students and faculty regarding the "
            "examination procedures for the 2022 examination cycle. "
            "NOTE: This version is superseded by the 2023 circular.\n\n"
            "Departments are advised to use the 2023 guidelines for current procedures.",
        ),
        (
            "1. Hall Ticket",
            "Hall tickets will be issued physically at the examination section counter "
            "from 20 October 2022. Students must collect their hall ticket in person.\n\n"
            "This procedure has been updated in 2023 to digital-only distribution.",
        ),
        (
            "2. Seating Arrangement",
            "Seating will be displayed on the noticeboard 72 hours before exams in 2022 "
            "procedures. Updated to 48-hour notice in 2023.\n\n"
            "Students are advised to refer to the 2023 circular for current timelines.",
        ),
        (
            "3. Prohibited Items",
            "Mobile phones and electronic devices are prohibited in the examination hall "
            "as per 2022 guidelines. Smartwatches were added to the prohibited list only "
            "from 2023 onwards.",
        ),
    ],
)

CSE_POLICY = (
    "CSE Department Academic Policy 2024",
    [
        (
            "1. Scope and Purpose",
            "This policy document governs academic operations within the Department of "
            "Computer Science and Engineering (CSE). It supplements the university-level "
            "handbook and takes effect from January 2024.\n\n"
            "All CSE faculty members, staff, and enrolled students are bound by this policy.",
        ),
        (
            "2. Course Registration",
            "Students must register for courses within the first five working days of each "
            "semester through the academic portal. Late registration attracts a penalty of "
            "Rs 500.\n\n"
            "A minimum of 15 credits and a maximum of 25 credits may be registered per semester "
            "unless approved by the Department Academic Committee (DAC).",
        ),
        (
            "3. Lab Regulations",
            "All CSE labs are accessible to students Monday through Saturday, 9 AM to 6 PM. "
            "After-hours access requires written approval from the lab coordinator.\n\n"
            "Students must sign the lab register on entry and exit. Damage to equipment due "
            "to negligence will be charged to the responsible student(s).",
        ),
        (
            "4. Project and Internship Policy",
            "Final year students are required to complete a capstone project worth 6 credits. "
            "Projects must be submitted by the 12th week of the final semester.\n\n"
            "Internships of a minimum 6-week duration may be credited as an elective if the "
            "internship report is approved by the faculty advisor and DAC.",
        ),
        (
            "5. Anti-Plagiarism Policy",
            "All assignments, projects, and theses are subject to plagiarism checks using "
            "approved software. A similarity index above 20% will result in assignment rejection.\n\n"
            "Repeat offences will be escalated to the Disciplinary Committee and may result "
            "in course failure or suspension.",
        ),
    ],
)

INTERNAL_STAFF_POLICY = (
    "Internal Staff Policy – University Administration 2024",
    [
        (
            "CONFIDENTIAL — FOR STAFF USE ONLY",
            "This document contains internal policies applicable to university staff members. "
            "It is not intended for student distribution and is marked access=internal.\n\n"
            "Unauthorized distribution of this document is a violation of university policy "
            "and may result in disciplinary action.",
        ),
        (
            "1. Staff Working Hours",
            "Administrative staff working hours are 9:00 AM to 5:30 PM, Monday to Friday. "
            "Flexible work arrangements may be approved by department heads on a case-by-case basis.\n\n"
            "All staff must swipe in and out using the biometric system. Manual attendance "
            "registers are maintained as a backup.",
        ),
        (
            "2. Salary and Allowances",
            "Salaries are disbursed on the last working day of each month via direct bank transfer. "
            "Medical allowance: Rs 15,000/year. Transport allowance: Rs 6,000/year.\n\n"
            "Staff who have completed 5+ years of service are eligible for the Long Service "
            "Increment of 5% of basic salary.",
        ),
        (
            "3. Leave Policy",
            "Casual Leave: 12 days/year. Medical Leave: 10 days/year. Earned Leave: 15 days/year. "
            "Leave without pay may be sanctioned for up to 30 days with the registrar's approval.\n\n"
            "Leave requests must be submitted at least 48 hours in advance through the HR portal, "
            "except in cases of medical emergency.",
        ),
        (
            "4. Code of Conduct",
            "All staff members are expected to maintain professional conduct at all times. "
            "Harassment, discrimination, or misconduct of any kind will be dealt with under "
            "the University Disciplinary Framework.\n\n"
            "Staff must not engage in any outside employment or consultancy without prior "
            "written approval from the university administration.",
        ),
        (
            "5. IT and Data Security",
            "Staff must not share their university login credentials with any other person. "
            "Sensitive student and financial data must be stored only on university-approved "
            "systems.\n\n"
            "Violation of data security policies may lead to termination of employment and "
            "legal proceedings under applicable data protection laws.",
        ),
    ],
)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def generate_all() -> None:
    print(f"\nGenerating sample documents in: {DATA_DIR.resolve()}\n")

    _write_pdf(DATA_DIR / "handbook_2024.pdf", *HANDBOOK_2024)
    _write_pdf(DATA_DIR / "circular_exam_rules_2023.pdf", *CIRCULAR_EXAM_2023)
    _write_pdf(DATA_DIR / "circular_exam_rules_2022.pdf", *CIRCULAR_EXAM_2022)
    _write_docx(DATA_DIR / "cse_department_policy.docx", *CSE_POLICY)
    _write_html(DATA_DIR / "internal_staff_policy.html", *INTERNAL_STAFF_POLICY)

    print(f"\nDone! {len(list(DATA_DIR.iterdir()))} files in {DATA_DIR}")


if __name__ == "__main__":
    generate_all()
